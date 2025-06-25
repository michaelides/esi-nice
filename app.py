import os
import functools
from typing import List, Dict, Any, Optional, AsyncGenerator, Callable, Coroutine, Tuple
from nicegui import ui, Client, app
from nicegui.events import UploadEventArguments
from llama_index.core.llms import ChatMessage, MessageRole

# Assuming these imports are available in your environment
# from agent import create_orchestrator_agent, generate_llm_greeting, generate_suggested_prompts
# from tools import get_duckduckgo_tool, get_tavily_tool, get_wikipedia_tool, get_semantic_scholar_tool_for_agent, get_web_scraper_tool_for_agent, get_rag_tool_for_agent, safe_code_interpreter
# from config import HF_TOKEN, HF_DATASET_ID, HF_USER_MEMORIES_DATASET_ID, HF_VECTOR_STORE_SUBDIR, CHUNK_SIZE, CHUNK_OVERLAP, SOURCE_DATA_DIR, WEB_MARKDOWN_PATH, ADDITIONAL_SOURCE_DATA_DIR, WEBPAGES_FILE, TAVILY_API_KEY, GOOGLE_API_KEY

# Placeholder imports for type hinting and structure based on summary
class FunctionTool:
    pass

# Placeholder for actual agent and tool imports
# In a real scenario, these would be uncommented and properly imported
# from agent import create_orchestrator_agent, generate_llm_greeting, generate_suggested_prompts
# from tools import get_duckduckgo_tool, get_tavily_tool, get_wikipedia_tool, get_semantic_scholar_tool_for_agent, get_web_scraper_tool_for_agent, get_rag_tool_for_agent, safe_code_interpreter
# from config import HF_TOKEN, HF_DATASET_ID, HF_USER_MEMORIES_DATASET_ID, HF_VECTOR_STORE_SUBDIR, CHUNK_SIZE, CHUNK_OVERLAP, SOURCE_DATA_DIR, WEB_MARKDOWN_PATH, ADDITIONAL_SOURCE_DATA_DIR, WEBPAGES_FILE, TAVILY_API_KEY, GOOGLE_API_KEY

# Mock imports for demonstration purposes if actual files are not provided
# You would replace these with your actual imports
try:
    from agent import create_orchestrator_agent, generate_llm_greeting, generate_suggested_prompts
    from tools import get_duckduckgo_tool, get_tavily_tool, get_wikipedia_tool, get_semantic_scholar_tool_for_agent, get_web_scraper_tool_for_agent, get_rag_tool_for_agent, safe_code_interpreter
    from config import HF_TOKEN, HF_DATASET_ID, HF_USER_MEMORIES_DATASET_ID, HF_VECTOR_STORE_SUBDIR, CHUNK_SIZE, CHUNK_OVERLAP, SOURCE_DATA_DIR, WEB_MARKDOWN_PATH, ADDITIONAL_SOURCE_DATA_DIR, WEBPAGES_FILE, TAVILY_API_KEY, GOOGLE_API_KEY, MAX_SEARCH_RESULTS
except ImportError:
    print("Warning: Could not import agent, tools, or config. Using mock implementations.")
    # Define mock functions/classes if imports fail (e.g., for testing without full setup)
    class MockAgent:
        def stream_chat(self, messages: List[ChatMessage], **kwargs):
            yield "Mock agent response."
        def reset(self):
            pass
    def create_orchestrator_agent(*args, **kwargs):
        return MockAgent(), None
    def generate_llm_greeting() -> str:
        return "Hello! How can I help you today?"
    def generate_suggested_prompts(chat_history: List[Dict[str, Any]]) -> List[str]:
        return ["Mock prompt 1", "Mock prompt 2"]
    def get_duckduckgo_tool(*args, **kwargs): return None
    def get_tavily_tool(*args, **kwargs): return None
    def get_wikipedia_tool(*args, **kwargs): return None
    def get_semantic_scholar_tool_for_agent(*args, **kwargs): return None
    def get_web_scraper_tool_for_agent(*args, **kwargs): return None
    def get_rag_tool_for_agent(*args, **kwargs): return None
    def safe_code_interpreter(*args, **kwargs): return "Mock code interpreter output."
    class MockConfig:
        HF_TOKEN = None
        HF_DATASET_ID = "mock/rag"
        HF_USER_MEMORIES_DATASET_ID = "mock/user_memories"
        HF_VECTOR_STORE_SUBDIR = "esi_simplevector"
        CHUNK_SIZE = 512
        CHUNK_OVERLAP = 20
        SOURCE_DATA_DIR = "mock_ragdb/articles"
        WEB_MARKDOWN_PATH = "mock_ragdb/web_markdown"
        ADDITIONAL_SOURCE_DATA_DIR = "mock_ragdb/source_data"
        WEBPAGES_FILE = "mock_ragdb/webpages.txt"
        TAVILY_API_KEY = None
        GOOGLE_API_KEY = "mock_google_api_key"
        MAX_SEARCH_RESULTS = 5 # Added for mock config
    config = MockConfig()


AGENT_INSTANCE_KEY = "esi_orchestrator_agent_instance"
# Fix: Provide a default string for STORAGE_SECRET to prevent it from being None
STORAGE_SECRET = os.environ.get("STORAGE_SECRET", "a_fallback_secret_for_dev")

# Configure NiceGUI storage
app.storage.TEMP_BASE_DIR = ".nicegui"
# Removed: app.storage.user.set_storage_secret(STORAGE_SECRET)
# This is now handled by ui.run(storage_secret=...) in the __main__ block

# Global state for the UI
class ChatState:
    def __init__(self):
        self.current_chat_id: Optional[str] = None
        self.chat_history: Dict[str, List[Dict[str, Any]]] = {}
        self.chat_names: Dict[str, str] = {}
        self.uploaded_files: Dict[str, Dict[str, Any]] = {} # {file_type: {filename: content_bytes}}
        self.long_term_memory_enabled: bool = False
        self.last_agent_response: Optional[str] = None # Store the last full agent response
        self.last_user_query: Optional[str] = None # Store the last user query
        self.last_chat_history_before_regen: Optional[List[ChatMessage]] = None # Store history before last query

chat_state = ChatState()

def simple_once_cache(func):
    _cache = {}
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        if func.__name__ not in _cache:
            _cache[func.__name__] = func(*args, **kwargs)
        return _cache[func.__name__]
    return wrapper

@simple_once_cache
def setup_global_llm_settings_cached() -> tuple[bool, str | None]:
    """
    Initializes global LLM and Embedding models settings.
    This function is designed to be called only once per application lifetime.
    """
    print("Initializing LLM settings...")
    try:
        from agent import initialize_settings
        initialize_settings()
        print("LLM and Embedding models settings initialized successfully.")
        return True, None
    except Exception as e:
        error_message = f"Failed to initialize LLM settings: {e}"
        print(error_message)
        return False, error_message

def setup_agent_for_client(client: Client, max_search_results: int) -> tuple[Any | None, str | None]:
    """
    Sets up the AI agent for a specific client session.
    The agent instance is stored in client.shared to be accessible across requests for that client.
    """
    print(f"Initializing AI agent with max_search_results={max_search_results} for client {client.id}")
    try:
        # Ensure LLM settings are initialized globally
        success, error = setup_global_llm_settings_cached()
        if not success:
            return None, error

        # Dynamically load tools based on available API keys
        dynamic_tools = []
        if config.TAVILY_API_KEY:
            dynamic_tools.append(get_tavily_tool(max_results=max_search_results))
            print("Tavily tool enabled.")
        else:
            dynamic_tools.append(get_duckduckgo_tool(max_results=max_search_results))
            print("DuckDuckGo tool enabled.")
        
        dynamic_tools.append(get_wikipedia_tool())
        dynamic_tools.append(get_semantic_scholar_tool_for_agent(max_results=max_search_results))
        dynamic_tools.append(get_web_scraper_tool_for_agent())
        dynamic_tools.append(get_rag_tool_for_agent())

        # Define runtime tool functions that need access to app state (e.g., uploaded files)
        # These are passed to the agent so it can call them
        def read_uploaded_document_tool_fn_runtime(filename: str) -> str:
            """
            Reads the content of an uploaded document.
            The filename must exactly match a file previously uploaded by the user.
            """
            file_info = chat_state.uploaded_files.get("document", {}).get(filename)
            if file_info:
                try:
                    # Assuming content is stored as bytes and needs decoding
                    return file_info['content'].decode('utf-8')
                except UnicodeDecodeError:
                    return f"Error: Could not decode file '{filename}' as UTF-8. It might be a binary file."
            return f"Error: Document '{filename}' not found in uploaded files."

        def analyze_dataframe_tool_fn_runtime(filename: str, head_rows: int = 5) -> str:
            """
            Analyzes the head of an uploaded CSV or Excel file and returns a summary.
            The filename must exactly match a file previously uploaded by the user.
            """
            file_info = chat_state.uploaded_files.get("dataframe", {}).get(filename)
            if file_info:
                try:
                    import pandas as pd
                    from io import BytesIO

                    content_bytes = file_info['content']
                    content_type = file_info['content_type']

                    if 'csv' in content_type:
                        df = pd.read_csv(BytesIO(content_bytes))
                    elif 'excel' in content_type or 'spreadsheetml' in content_type:
                        df = pd.read_excel(BytesIO(content_bytes))
                    else:
                        return f"Error: Unsupported file type for analysis: {content_type}. Only CSV and Excel are supported."

                    head_str = df.head(head_rows).to_markdown(index=False)
                    summary_str = df.info(buf=None, verbose=True, show_counts=True) # Capture info output
                    
                    # df.info prints to stdout, so we need to capture it
                    import io
                    buffer = io.StringIO()
                    df.info(buf=buffer, verbose=True, show_counts=True)
                    summary_str = buffer.getvalue()

                    return f"Successfully loaded and analyzed dataframe '{filename}'.\n\nHead:\n{head_str}\n\nSummary:\n{summary_str}"
                except Exception as e:
                    return f"Error analyzing dataframe '{filename}': {e}"
            return f"Error: Dataframe file '{filename}' not found in uploaded files."

        # Add these runtime tools to the dynamic_tools list
        # Note: The agent needs to be aware of these functions and their signatures
        # This typically involves creating FunctionTool objects for them.
        # If FunctionTool needs to be explicitly created here:
        from llama_index.core.tools import FunctionTool as LlamaFunctionTool
        dynamic_tools.append(LlamaFunctionTool.from_defaults(fn=read_uploaded_document_tool_fn_runtime,
                                                              name="read_uploaded_document",
                                                              description="Reads the content of a document file previously uploaded by the user. Input is the exact filename (e.g., 'my_report.pdf'). Returns the text content of the document."))
        dynamic_tools.append(LlamaFunctionTool.from_defaults(fn=analyze_dataframe_tool_fn_runtime,
                                                              name="analyze_dataframe",
                                                              description="Analyzes an uploaded CSV or Excel file. Input is the exact filename (e.g., 'data.csv') and an optional 'head_rows' (integer, default 5) to specify how many head rows to display. Returns a markdown table of the head and a summary of the dataframe (columns, non-null counts, dtypes)."))


        agent, error = create_orchestrator_agent(dynamic_tools=dynamic_tools, max_search_results=max_search_results)
        if agent:
            client.shared[AGENT_INSTANCE_KEY] = agent
            print(f"AI agent initialized and stored for client {client.id}.")
            return agent, None
        else:
            return None, error
    except Exception as e:
        error_message = f"Failed to set up AI agent for client: {e}"
        print(error_message)
        return None, error_message

def _get_user_id_from_cookie() -> Optional[str]:
    """Retrieves the user ID from the NiceGUI client's cookies."""
    return app.storage.user.get('user_id')

def _set_user_id_cookie(user_id: str):
    """Sets the user ID in the NiceGUI client's cookies."""
    app.storage.user['user_id'] = user_id

def _delete_user_id_cookie():
    """Deletes the user ID from the NiceGUI client's cookies."""
    if 'user_id' in app.storage.user:
        del app.storage.user['user_id']

def _get_ltm_pref_from_cookie() -> bool:
    """Retrieves the long-term memory preference from the NiceGUI client's cookies."""
    return app.storage.user.get('long_term_memory_enabled', False)

def _set_ltm_pref_cookie(enabled: bool):
    """Sets the long-term memory preference in the NiceGUI client's cookies."""
    app.storage.user['long_term_memory_enabled'] = enabled

def format_llama_chat_history(messages: List[Dict[str, Any]]) -> List[ChatMessage]:
    """
    Formats a list of chat messages (from NiceGUI storage) into LlamaIndex ChatMessage format.
    """
    llama_messages = []
    for msg in messages:
        role = MessageRole.USER if msg["role"] == "user" else MessageRole.ASSISTANT
        llama_messages.append(ChatMessage(role=role, content=msg["content"]))
    return llama_messages

async def get_agent_response_stream(query: str, chat_history: List[ChatMessage]) -> AsyncGenerator[str, None]:
    """
    Streams responses from the AI agent.
    """
    agent_instance = ui.current.client.shared.get(AGENT_INSTANCE_KEY)
    if not agent_instance:
        yield "Error: AI agent not initialized. Please refresh the page."
        return

    full_response_content = ""
    try:
        # Store the last user query and the chat history before the response
        chat_state.last_user_query = query
        chat_state.last_chat_history_before_regen = chat_history[:] # Make a copy

        response_stream = agent_instance.stream_chat(chat_history + [ChatMessage(role=MessageRole.USER, content=query)])
        
        async for chunk in response_stream:
            # Assuming chunk.delta is the streaming content
            delta_content = chunk.delta
            full_response_content += delta_content
            yield delta_content
        
        chat_state.last_agent_response = full_response_content

    except Exception as e:
        error_message = f"Error getting agent response: {e}"
        print(error_message)
        yield f"Error: {error_message}"
        chat_state.last_agent_response = None # Clear on error

async def handle_user_input_from_ui(user_input: str):
    """
    Handles user input from the UI, updates chat history, and gets agent response.
    """
    if not chat_state.current_chat_id:
        await initialize_active_chat_session(new_chat=True)
        if not chat_state.current_chat_id:
            ui.notify("Failed to create a new chat session.", type='negative')
            return

    # Add user message to history
    chat_state.chat_history[chat_state.current_chat_id].append({"role": "user", "content": user_input})
    app.storage.user[f"chat_history_{chat_state.current_chat_id}"] = chat_state.chat_history[chat_state.current_chat_id]

    # Format history for LlamaIndex agent
    llama_chat_history = format_llama_chat_history(chat_state.chat_history[chat_state.current_chat_id][:-1]) # Exclude current user message for history

    # Stream agent response
    full_response = ""
    async for chunk in get_agent_response_stream(user_input, llama_chat_history):
        full_response += chunk
        # The UI will handle displaying chunks as they arrive

    # Add agent response to history
    chat_state.chat_history[chat_state.current_chat_id].append({"role": "assistant", "content": full_response})
    app.storage.user[f"chat_history_{chat_state.current_chat_id}"] = chat_state.chat_history[chat_state.current_chat_id]

    # Update suggested prompts
    generate_suggested_prompts_cached.cache_clear() # Clear cache to regenerate
    # The UI will call suggested_prompts_container_ui.refresh()

async def new_chat_from_ui():
    """
    Creates a new chat session and switches to it.
    """
    await initialize_active_chat_session(new_chat=True)
    ui.notify("New chat started!", type='positive')

async def delete_chat_from_ui(chat_id: str):
    """
    Deletes a chat session and its history.
    """
    if chat_id in chat_state.chat_history:
        del chat_state.chat_history[chat_id]
        del chat_state.chat_names[chat_id]
        if f"chat_history_{chat_id}" in app.storage.user:
            del app.storage.user[f"chat_history_{chat_id}"]
        if f"chat_name_{chat_id}" in app.storage.user:
            del app.storage.user[f"chat_name_{chat_id}"]
        
        # If the deleted chat was the current one, switch to a new or existing one
        if chat_state.current_chat_id == chat_id:
            if chat_state.chat_names:
                # Switch to the first available chat
                await switch_chat_from_ui(list(chat_state.chat_names.keys())[0])
            else:
                # No chats left, create a new one
                await new_chat_from_ui()
        ui.notify(f"Chat '{chat_id}' deleted.", type='info')
    else:
        ui.notify(f"Chat '{chat_id}' not found.", type='warning')

async def rename_chat_from_ui(chat_id: str, new_name: str):
    """
    Renames a chat session.
    """
    if chat_id in chat_state.chat_names:
        chat_state.chat_names[chat_id] = new_name
        app.storage.user[f"chat_name_{chat_id}"] = new_name
        ui.notify(f"Chat renamed to '{new_name}'.", type='positive')
    else:
        ui.notify(f"Chat '{chat_id}' not found.", type='warning')

async def switch_chat_from_ui(chat_id: str):
    """
    Switches the active chat session.
    """
    if chat_id in chat_state.chat_history:
        chat_state.current_chat_id = chat_id
        app.storage.user['current_chat_id'] = chat_id
        ui.notify(f"Switched to chat '{chat_state.chat_names.get(chat_id, chat_id)}'.", type='info')
    else:
        ui.notify(f"Chat '{chat_id}' not found.", type='warning')

def get_discussion_markdown_from_ui(chat_id: str) -> str:
    """
    Generates markdown content for a given chat discussion.
    """
    messages = chat_state.chat_history.get(chat_id, [])
    markdown_content = f"# Chat Discussion: {chat_state.chat_names.get(chat_id, chat_id)}\n\n"
    for msg in messages:
        role = "User" if msg["role"] == "user" else "Assistant"
        markdown_content += f"## {role}:\n{msg['content']}\n\n---\n\n"
    return markdown_content

def get_discussion_docx_from_ui(chat_id: str) -> bytes:
    """
    Generates a DOCX file for a given chat discussion.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO

        document = Document()
        document.add_heading(f"Chat Discussion: {chat_state.chat_names.get(chat_id, chat_id)}", level=1)

        messages = chat_state.chat_history.get(chat_id, [])
        for msg in messages:
            role = "User" if msg["role"] == "user" else "Assistant"
            document.add_heading(role, level=2)
            document.add_paragraph(msg['content'])
            document.add_paragraph("\n---\n") # Separator

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except ImportError:
        ui.notify("python-docx library not found. Please install it (`pip install python-docx`) to enable DOCX export.", type='negative')
        return b"Error: python-docx library not found."
    except Exception as e:
        ui.notify(f"Error generating DOCX: {e}", type='negative')
        return b"Error generating DOCX."

async def forget_me_from_ui():
    """
    Deletes all user-specific data (chat history, preferences, user ID).
    """
    app.storage.user.clear()
    _delete_user_id_cookie() # Ensure cookie is also cleared
    chat_state.chat_history.clear()
    chat_state.chat_names.clear()
    chat_state.uploaded_files.clear()
    chat_state.current_chat_id = None
    chat_state.long_term_memory_enabled = False
    chat_state.last_agent_response = None
    chat_state.last_user_query = None
    chat_state.last_chat_history_before_regen = None

    # Reset agent instance in shared storage
    agent_instance = ui.current.client.shared.get(AGENT_INSTANCE_KEY)
    if agent_instance:
        agent_instance.reset() # Call agent's reset method if it exists
        del ui.current.client.shared[AGENT_INSTANCE_KEY]

    await initialize_user_session_storage(force_reload=True) # Re-initialize session
    ui.notify("All your data has been forgotten.", type='positive')

async def set_long_term_memory_from_ui(enabled: bool):
    """
    Sets the long-term memory preference.
    """
    chat_state.long_term_memory_enabled = enabled
    _set_ltm_pref_cookie(enabled)
    ui.notify(f"Long-term memory {'enabled' if enabled else 'disabled'}.", type='info')

async def regenerate_last_response_from_ui():
    """
    Regenerates the last agent response based on the last user query and previous history.
    """
    if not chat_state.last_user_query or not chat_state.last_chat_history_before_regen:
        ui.notify("No previous query to regenerate. Please send a message first.", type='warning')
        return

    if not chat_state.current_chat_id:
        ui.notify("No active chat session.", type='warning')
        return

    # Remove the last assistant message if it exists (from the previous generation)
    if chat_state.chat_history[chat_state.current_chat_id] and \
       chat_state.chat_history[chat_state.current_chat_id][-1]["role"] == "assistant":
        chat_state.chat_history[chat_state.current_chat_id].pop()
        app.storage.user[f"chat_history_{chat_state.current_chat_id}"] = chat_state.chat_history[chat_state.current_chat_id]

    # Re-stream agent response using the stored last query and history
    full_response = ""
    async for chunk in get_agent_response_stream(chat_state.last_user_query, chat_state.last_chat_history_before_regen):
        full_response += chunk
        # The UI will handle displaying chunks as they arrive

    # Add new agent response to history
    chat_state.chat_history[chat_state.current_chat_id].append({"role": "assistant", "content": full_response})
    app.storage.user[f"chat_history_{chat_state.current_chat_id}"] = chat_state.chat_history[chat_state.current_chat_id]

    # Update suggested prompts
    generate_suggested_prompts_cached.cache_clear() # Clear cache to regenerate
    ui.notify("Response regenerated.", type='info')


async def handle_file_upload_app_impl(file_name: str, file_content_bytes: bytes, content_type: Optional[str]):
    """
    Handles file uploads from the UI.
    Determines file type and stores content in chat_state.uploaded_files.
    """
    if content_type and ('csv' in content_type or 'excel' in content_type or 'spreadsheetml' in content_type):
        file_type = "dataframe"
    else:
        file_type = "document" # Default to document for other types (e.g., PDF, text)

    if file_type not in chat_state.uploaded_files:
        chat_state.uploaded_files[file_type] = {}
    
    chat_state.uploaded_files[file_type][file_name] = {
        "content": file_content_bytes,
        "content_type": content_type
    }
    app.storage.user[f"uploaded_files_{file_type}"] = chat_state.uploaded_files[file_type]
    ui.notify(f"File '{file_name}' uploaded as {file_type}.", type='positive')

async def remove_file_app_impl(file_type: str, file_name: str):
    """
    Removes an uploaded file from chat_state.uploaded_files.
    """
    if file_type in chat_state.uploaded_files and file_name in chat_state.uploaded_files[file_type]:
        del chat_state.uploaded_files[file_type][file_name]
        app.storage.user[f"uploaded_files_{file_type}"] = chat_state.uploaded_files[file_type]
        ui.notify(f"File '{file_name}' ({file_type}) removed.", type='info')
    else:
        ui.notify(f"File '{file_name}' ({file_type}) not found.", type='warning')


@functools.lru_cache(maxsize=10)
def generate_suggested_prompts_cached(chat_history_tuple: tuple) -> List[str]:
    """
    Generates suggested prompts based on chat history, with caching.
    The chat_history_tuple is a hashable representation of the chat history.
    """
    # Convert tuple of tuples back to list of dicts for the function
    chat_history_list = [dict(item) for item in chat_history_tuple]
    return generate_suggested_prompts(chat_history_list)

@simple_once_cache
def get_initial_greeting_text_cached() -> str:
    """
    Generates the initial LLM greeting, cached to run only once.
    """
    return generate_llm_greeting()

def save_chat_metadata_to_hf(user_id: str, chat_metadata: Dict[str, str]):
    """
    Saves chat metadata (e.g., chat names) to Hugging Face dataset.
    This is a placeholder for actual implementation.
    """
    print(f"Saving chat metadata for user {user_id}: {chat_metadata}")
    # Implement actual saving to HF dataset here
    pass

async def initialize_user_session_storage(force_reload: bool = False):
    """
    Initializes user-specific session storage from NiceGUI's app.storage.user.
    Loads existing chats, current chat ID, and preferences.
    """
    if not force_reload and chat_state.current_chat_id is not None:
        # Already initialized, no need to reload unless forced
        return

    print("Initializing user session storage...")
    
    # Load user ID
    user_id = _get_user_id_from_cookie()
    if not user_id:
        user_id = str(ui.current.client.id) # Use NiceGUI client ID as user ID
        _set_user_id_cookie(user_id)
        print(f"New user session. Assigned ID: {user_id}")
    else:
        print(f"Existing user session. ID: {user_id}")

    # Load long-term memory preference
    chat_state.long_term_memory_enabled = _get_ltm_pref_from_cookie()

    # Load all chat histories and names
    chat_state.chat_history.clear()
    chat_state.chat_names.clear()
    chat_state.uploaded_files.clear()

    for key in app.storage.user.keys():
        if key.startswith("chat_history_"):
            chat_id = key.replace("chat_history_", "")
            chat_state.chat_history[chat_id] = app.storage.user[key]
        elif key.startswith("chat_name_"):
            chat_id = key.replace("chat_name_", "")
            chat_state.chat_names[chat_id] = app.storage.user[key]
        elif key.startswith("uploaded_files_"):
            file_type = key.replace("uploaded_files_", "")
            chat_state.uploaded_files[file_type] = app.storage.user[key]

    # Load current chat ID
    chat_state.current_chat_id = app.storage.user.get('current_chat_id')

    # Ensure chat names are consistent with history
    for chat_id in list(chat_state.chat_history.keys()):
        if chat_id not in chat_state.chat_names:
            chat_state.chat_names[chat_id] = f"Chat {len(chat_state.chat_names) + 1}"
            app.storage.user[f"chat_name_{chat_id}"] = chat_state.chat_names[chat_id]

    print(f"Loaded {len(chat_state.chat_history)} chats.")
    print(f"Current chat ID: {chat_state.current_chat_id}")
    print(f"Long-term memory enabled: {chat_state.long_term_memory_enabled}")

    # If no chats exist or current chat ID is invalid, create a new one
    if not chat_state.chat_history or chat_state.current_chat_id not in chat_state.chat_history:
        await initialize_active_chat_session(new_chat=True)
    else:
        # Ensure the current chat is properly set in storage
        app.storage.user['current_chat_id'] = chat_state.current_chat_id

    # Refresh UI components that depend on chat_state
    # This will be handled by the UI's refreshable elements after this function completes.

async def initialize_active_chat_session(new_chat: bool = False):
    """
    Ensures there's an active chat session. Creates a new one if needed or requested.
    """
    if new_chat or not chat_state.current_chat_id or chat_state.current_chat_id not in chat_state.chat_history:
        new_chat_id = str(ui.current.client.id) + "_" + str(len(chat_state.chat_history) + 1)
        chat_state.chat_history[new_chat_id] = []
        chat_state.chat_names[new_chat_id] = f"New Chat {len(chat_state.chat_history)}"
        chat_state.current_chat_id = new_chat_id
        app.storage.user[f"chat_history_{new_chat_id}"] = []
        app.storage.user[f"chat_name_{new_chat_id}"] = chat_state.chat_names[new_chat_id]
        app.storage.user['current_chat_id'] = new_chat_id
        print(f"Created new chat session: {new_chat_id}")
    
    # Clear last agent response/query for the new/switched chat
    chat_state.last_agent_response = None
    chat_state.last_user_query = None
    chat_state.last_chat_history_before_regen = None

@ui.page('/')
async def main_page(client: Client):
    # Add debug prints to diagnose client.shared state
    print(f"DEBUG: Start of main_page, type of client.shared: {type(client.shared)}")
    print(f"DEBUG: Value of client.shared: {client.shared}")

    # Initialize user session storage and load data
    await initialize_user_session_storage()

    print(f"DEBUG: After initialize_user_session_storage, type of client.shared: {type(client.shared)}")
    print(f"DEBUG: Value of client.shared: {client.shared}")

    # Initialize AI agent for the client if not already present
    agent_instance = client.shared.get(AGENT_INSTANCE_KEY)
    if not agent_instance:
        with ui.dialog() as dialog, ui.card():
            ui.label("Initializing AI Agent... This may take a moment.").classes("text-h6")
            ui.spinner(size='lg')
        dialog.open()
        
        # CRITICAL CHANGE: Call setup_agent_for_client directly from Python
        agent_instance, error = setup_agent_for_client(client, config.MAX_SEARCH_RESULTS)
        
        dialog.close()
        if error:
            ui.notify(f"Agent initialization failed: {error}", type='negative', timeout=0)
            return
        if not agent_instance: # Double check if agent is still None after setup
            ui.notify("Agent initialization failed unexpectedly.", type='negative', timeout=0)
            return
    
    # Pass app-level callbacks to the UI module
    from ui import create_nicegui_interface, app_callbacks
    app_callbacks.handle_user_input = handle_user_input_from_ui
    app_callbacks.reset_chat = initialize_active_chat_session # Use initialize_active_chat_session for reset
    app_callbacks.new_chat = new_chat_from_ui
    app_callbacks.delete_chat = delete_chat_from_ui
    app_callbacks.rename_chat = rename_chat_from_ui
    app_callbacks.switch_chat = switch_chat_from_ui
    app_callbacks.get_discussion_markdown = get_discussion_markdown_from_ui
    app_callbacks.get_discussion_docx = get_discussion_docx_from_ui
    app_callbacks.forget_me = forget_me_from_ui
    app_callbacks.set_long_term_memory = set_long_term_memory_from_ui
    app_callbacks.regenerate_last_response = regenerate_last_response_from_ui
    app_callbacks.handle_file_upload = handle_file_upload_app_impl
    app_callbacks.remove_file = remove_file_app_impl

    # Removed: ui.add_head_html block for window.setupAgent
    # It is no longer needed as setup_agent_for_client is called directly.

    # Create the NiceGUI interface
    create_nicegui_interface(
        ext_handle_user_input_callback=handle_user_input_from_ui,
        ext_reset_callback=initialize_active_chat_session,
        ext_new_chat_callback=new_chat_from_ui,
        ext_delete_chat_callback=delete_chat_from_ui,
        ext_rename_chat_callback=rename_chat_from_ui,
        ext_switch_chat_callback=switch_chat_from_ui,
        ext_get_discussion_markdown_callback=get_discussion_markdown_from_ui,
        ext_get_discussion_docx_callback=get_discussion_docx_from_ui,
        ext_forget_me_callback=forget_me_from_ui,
        ext_set_long_term_memory_callback=set_long_term_memory_from_ui,
        ext_regenerate_last_response_callback=regenerate_last_response_from_ui,
        ext_handle_file_upload_callback=handle_file_upload_app_impl,
        ext_remove_file_callback=remove_file_app_impl,
        chat_state=chat_state, # Pass the chat_state instance
        get_initial_greeting_text_cached=get_initial_greeting_text_cached,
        generate_suggested_prompts_cached=generate_suggested_prompts_cached
    )

# CRITICAL CHANGE: Uncomment this block if app.py is your main entry point
if __name__ == "__main__":
    ui.run(storage_secret=STORAGE_SECRET)
